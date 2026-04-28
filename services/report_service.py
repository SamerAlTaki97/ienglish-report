import io
import json
import os
import re
import base64
import shutil
import tempfile
import time
from datetime import datetime, timezone
from zoneinfo import ZoneInfo
from urllib.parse import urljoin

from flask import current_app, has_request_context, render_template, request
import models
from database import REPORT_STATUSES

from . import ServiceError
from . import ai_client
from .user_service import user_has_role

ALLOWED_TRANSITIONS = {
    "draft": {"pending_operation"},
    "pending_operation": {"approved", "rejected"},
    "approved": {"delivered"},
    "delivered": set(),
    "rejected": set(),
}

DATASET_FILE = "dataset.json"
COURSE_TYPES = {
    "general_english": "General English",
    "ielts": "IELTS",
    "sat": "SAT",
    "other": "Other",
}
CLASS_TYPES = {"Group Adult", "VIP Adult", "Group Kids", "VIP Kid"}
MR_NAMES = {"hussam", "fouad", "mazen", "zain", "amin", "karam", "omar"}
DIRECT_SALES_VALUE = "direct"
DIRECT_SALES_LABEL = "Direct"
DISPLAY_TIMEZONE = ZoneInfo("Asia/Dubai")


def generate_report(data, actor):
    if not user_has_role(actor, "teacher", "admin", "manager"):
        raise ServiceError("Only teachers can create reports", 403)

    submission = _normalize_submission(data, actor)
    ai_payload = _generate_ai_report(submission)
    report_id = _persist_report(submission, ai_payload, actor)
    _append_dataset_entry(submission, ai_payload)

    report = models.get_report_detail(report_id, include_json=True)
    report["report_json"]["report_id"] = report_id
    report["report_json"]["status"] = report["status"]
    report["report_json"]["pdf_url"] = report["pdf_path"]
    return report


def update_report(report_id, data, actor):
    report = get_report_for_user(report_id, actor, include_json=True)
    if report["status"] != "draft":
        raise ServiceError("Only draft reports can be edited", 400)
    if not (user_has_role(actor, "admin", "manager") or actor["id"] == report["teacher_id"]):
        raise ServiceError("You cannot edit this report", 403)

    submission = _normalize_submission(data, actor, existing_report=report)
    ai_payload = _generate_ai_report(submission)
    report_payload = _compose_report_payload(submission, ai_payload)

    if submission["student"].get("phone"):
        models.upsert_student(
            phone=submission["student"]["phone"],
            name=submission["student"]["name"],
            branch=submission["student"].get("branch"),
            email=submission["student"].get("email"),
            sales_id=submission["student"].get("sales_id"),
        )
    models.update_report_content(
        report_id=report_id,
        report_type=submission["report_type"],
        level=submission["student"].get("level"),
        course=submission["student"].get("course"),
        report_json=report_payload,
    )
    _store_pdf(
        report_id,
        generate_pdf_bytes(report_payload),
        submission["student"],
        submission["report_type"],
    )
    models.add_report_tracking(report_id, "draft", actor["id"], "Teacher updated draft")

    updated = models.get_report_detail(report_id, include_json=True)
    updated["report_json"]["report_id"] = report_id
    updated["report_json"]["status"] = updated["status"]
    updated["report_json"]["pdf_url"] = updated["pdf_path"]
    return updated


def submit_report(report_id, actor):
    report = get_report_for_user(report_id, actor, include_json=False)
    if not (user_has_role(actor, "admin", "manager") or actor["id"] == report["teacher_id"]):
        raise ServiceError("You cannot submit this report", 403)
    return _transition_report(
        report_id,
        actor,
        "pending_operation",
        "Teacher submitted report to operation",
    )


def approve_report(report_id, actor):
    if not user_has_role(actor, "operation", "admin", "manager"):
        raise ServiceError("Only operation can approve reports", 403)
    return _transition_report(
        report_id,
        actor,
        "approved",
        "Operation approved report",
    )


def reject_report(report_id, actor):
    if not user_has_role(actor, "operation", "admin", "manager"):
        raise ServiceError("Only operation can reject reports", 403)
    return _transition_report(
        report_id,
        actor,
        "rejected",
        "Operation marked report as not accepted",
    )


def mark_report_delivered(report_id, actor, note):
    if not user_has_role(actor, "operation", "admin", "manager"):
        raise ServiceError("Only operation can mark reports as delivered", 403)
    return _transition_report(report_id, actor, "delivered", note)


def delete_report(report_id, actor):
    if not user_has_role(actor, "operation", "admin", "manager"):
        raise ServiceError("Only operation can delete reports", 403)
    report = get_report_for_user(report_id, actor, include_json=False)
    if report["status"] not in {"pending_operation", "rejected"}:
        raise ServiceError("Only pending or not accepted reports can be deleted", 400)
    if not models.delete_report(report_id):
        raise ServiceError("Report not found", 404)
    return {"id": report_id, "deleted": True}


def get_report_for_user(report_id, actor, include_json=True):
    report = models.get_report_detail(report_id, include_json=include_json)
    if not report:
        raise ServiceError("Report not found", 404)
    if not _can_view_report(actor, report):
        raise ServiceError("Forbidden", 403)
    return report


def get_public_report_pdf(report_id):
    report = models.get_report_detail(report_id, include_json=False)
    if not report:
        raise ServiceError("Report not found", 404)

    document = models.get_report_document(report_id)
    if not document:
        raise ServiceError("PDF not found", 404)

    return {
        "report": report,
        "document": document,
    }


def refresh_report_pdf(report_id):
    report = models.get_report_detail(report_id, include_json=True)
    if not report:
        raise ServiceError("Report not found", 404)

    report_payload = report.get("report_json") or {}
    student_payload = dict(report_payload.get("student") or {})
    student_payload["phone"] = _effective_student_phone(report)
    student_payload["email"] = report.get("student_email")
    student_payload["name"] = student_payload.get("name") or report.get("student_name")
    student_payload["branch"] = student_payload.get("branch") or report.get("student_branch")
    report_payload["student"] = student_payload

    pdf_bytes = generate_pdf_bytes(report_payload)
    _store_pdf(
        report_id,
        pdf_bytes,
        student_payload,
        report.get("report_type"),
    )
    models.update_report_content(report_id=report_id, report_json=report_payload)
    return models.get_report_detail(report_id, include_json=False)


def list_reports_for_user(actor, phone=None, status=None, teacher_id=None, sales_id=None, branch=None, workflow_only=False):
    normalized_teacher_id = _normalize_id(teacher_id)
    normalized_sales_id = _normalize_id(sales_id)
    if actor.get("role") == "superadmin":
        reports = models.list_reports(phone=phone, status=status, teacher_id=normalized_teacher_id, sales_id=normalized_sales_id, branch=branch)
    elif user_has_role(actor, "manager"):
        reports = models.list_reports(
            phone=phone,
            status=status,
            teacher_id=normalized_teacher_id,
            sales_id=normalized_sales_id,
            branch=branch,
        )
    elif user_has_role(actor, "admin", "operation", "sales_admin"):
        reports = models.list_reports(
            phone=phone,
            status=status,
            teacher_id=normalized_teacher_id,
            sales_id=normalized_sales_id,
            branch=actor.get("branch"),
        )
    elif user_has_role(actor, "teacher"):
        reports = models.list_reports(phone=phone, status=status, teacher_id=actor["id"])
    elif user_has_role(actor, "sales"):
        reports = models.list_reports(phone=phone, status=status, sales_id=actor["id"])
    else:
        raise ServiceError("Forbidden", 403)

    if workflow_only:
        reports = [
            row for row in reports
            if row.get("status") in {"pending_operation", "approved", "delivered", "rejected"}
        ]

    return [_serialize_listing(row) for row in reports]


def search_students(phone_query, actor):
    if not user_has_role(actor, "teacher", "operation", "admin", "sales", "sales_admin", "manager"):
        raise ServiceError("Forbidden", 403)
    return models.search_students_by_phone(phone_query)


def generate_word_document(data):
    return generate_docx_bytes(data)


def generate_chart(data):
    return create_chart_bytes(data)


def render_report_html(data, preview_mode=False, report_meta=None, template_context=None):
    context = build_report_document_context(data, report_meta=report_meta, preview_mode=preview_mode)
    if template_context:
        context.update(template_context)
    return render_template("report_document.html", **context)


def build_report_download_name(report):
    report_payload = report.get("report_json") or {}
    if isinstance(report_payload, str):
        try:
            report_payload = json.loads(report_payload)
        except json.JSONDecodeError:
            report_payload = {}
    student_payload = dict(report_payload.get("student") or {})
    base_name = student_payload.get("name") or report.get("student_name") or "student"
    normalized = re.sub(r"[^0-9A-Za-z]+", " ", str(base_name).strip().lower())
    normalized = re.sub(r"\s+", " ", normalized).strip() or "student"
    return f"{normalized} report.pdf"


def build_report_document_context(data, report_meta=None, preview_mode=False):
    student = dict(data.get("student", {}) or {})
    student["name"] = _format_name_words(student.get("name"))
    if student.get("teacher"):
        student["teacher"] = _display_person_name(student.get("teacher"))
    skill_scores = _collect_skill_scores(data)
    snapshot_scores = _collect_snapshot_scores(data)
    performance = data.get("performance") or {}
    overall_score = _compute_average(list(skill_scores.values()) + list(snapshot_scores.values()))
    strongest_skill, weakest_skill = _get_strength_markers(skill_scores)
    sections = {str(item.get("title") or "").strip().upper(): str(item.get("content") or "").strip() for item in data.get("sections", [])}

    return {
        "preview_mode": preview_mode,
        "report_type_label": _format_report_type(data.get("report_type")),
        "issue_date": datetime.now(DISPLAY_TIMEZONE).strftime("%d %B %Y"),
        "student": student,
        "exam_score": performance.get("exam_score"),
        "report_meta": report_meta or {},
        "overall_score": overall_score,
        "performance_band": _score_band(overall_score),
        "strongest_skill": strongest_skill,
        "weakest_skill": weakest_skill,
        "snapshot_scores": snapshot_scores,
        "skill_scores": skill_scores,
        "executive_summary": _build_executive_summary(data, student, overall_score, strongest_skill, weakest_skill),
        "sections_map": sections,
        "analysis_sections": [
            {"title": "Introduction", "key": "INTRODUCTION", "content": sections.get("INTRODUCTION", "")},
            {"title": "Attendance & Performance", "key": "ATTENDANCE & PERFORMANCE", "content": sections.get("ATTENDANCE & PERFORMANCE", "")},
            {"title": "Communication Skills", "key": "COMMUNICATION SKILLS", "content": sections.get("COMMUNICATION SKILLS", "")},
            {"title": "Vocabulary Development", "key": "VOCABULARY DEVELOPMENT", "content": sections.get("VOCABULARY DEVELOPMENT", "")},
            {"title": "Grammar Analysis", "key": "GRAMMAR ANALYSIS", "content": sections.get("GRAMMAR ANALYSIS", "")},
            {"title": "Reading & Writing", "key": "READING & WRITING", "content": sections.get("READING & WRITING", "")},
        ],
        "next_steps": _derive_next_steps(data, strongest_skill, weakest_skill),
        "skills_chart_uri": _to_data_uri(create_chart_bytes(data), "image/png"),
        "snapshot_chart_uri": _to_data_uri(create_progress_chart_bytes(data), "image/png"),
        "logo_uri": _load_logo_data_uri(),
    }


def update_student_contact(report_id, actor, phone=None, email=None, sales_id=None, refresh_pdf=True):
    if not user_has_role(actor, "operation", "admin", "manager"):
        raise ServiceError("Only operation can update student contact details", 403)

    report = models.get_report_detail(report_id, include_json=False)
    if not report:
        raise ServiceError("Report not found", 404)

    raw_phone = str(phone).strip() if isinstance(phone, str) else None
    normalized_phone = raw_phone or _effective_student_phone(report) or None
    normalized_email = email.strip() if isinstance(email, str) else email
    sales_id_provided = sales_id is not None
    normalized_sales_id, sales_mode, sales_name = _normalize_sales_selection(sales_id)

    if not normalized_phone and normalized_email in ("", None) and normalized_sales_id is None and not sales_id_provided:
        raise ServiceError("Student phone, email, or educational consultant is required", 400)

    try:
        models.update_student_contact_for_report(
            report_id,
            phone=raw_phone,
            email=normalized_email,
            sales_id=normalized_sales_id,
            sales_id_provided=sales_id_provided,
        )
    except ValueError as exc:
        raise ServiceError(str(exc), 409) from exc

    updated_report = models.get_report_detail(report_id, include_json=True)
    report_payload = updated_report.get("report_json") or {}
    student_payload = dict(report_payload.get("student") or {})
    student_payload["phone"] = normalized_phone or _effective_student_phone(updated_report)
    student_payload["email"] = normalized_email if normalized_email not in ("", None) else updated_report.get("student_email")
    if sales_id_provided:
        student_payload["sales_id"] = normalized_sales_id
        student_payload["sales_mode"] = sales_mode
        if sales_name:
            student_payload["sales_name"] = sales_name
        else:
            student_payload.pop("sales_name", None)
    else:
        student_payload["sales_id"] = updated_report.get("sales_id")
        student_payload["sales_mode"] = updated_report.get("sales_mode") or student_payload.get("sales_mode")
    report_payload["student"] = student_payload
    models.update_report_content(report_id=report_id, report_json=report_payload)
    if refresh_pdf:
        try:
            _store_pdf(
                report_id,
                generate_pdf_bytes(report_payload),
                student_payload,
                updated_report["report_type"],
            )
        except ServiceError as exc:
            current_app.logger.warning("PDF refresh skipped for report %s: %s", report_id, exc.message)

    models.add_report_tracking(
        report_id,
        updated_report["status"],
        actor["id"],
        "Operation updated student contact details and consultant assignment",
    )
    return models.get_report_detail(report_id, include_json=True)


def _persist_report(submission, ai_payload, actor):
    student = submission["student"]
    report_payload = _compose_report_payload(submission, ai_payload)
    if student.get("phone"):
        models.upsert_student(
            phone=student["phone"],
            name=student["name"],
            branch=student.get("branch"),
            email=student.get("email"),
            sales_id=student.get("sales_id"),
        )

    report_id = models.create_report(
        student_id=student.get("phone"),
        teacher_id=actor["id"],
        report_type=submission["report_type"],
        level=student.get("level"),
        course=student.get("course"),
        report_json=report_payload,
        status="draft",
        created_by=actor["id"],
    )

    models.add_report_tracking(report_id, "draft", actor["id"], "Teacher created report")
    try:
        _store_pdf(
            report_id,
            generate_pdf_bytes(report_payload),
            student,
            submission["report_type"],
        )
    except ServiceError as exc:
        current_app.logger.warning("PDF generation skipped for report %s: %s", report_id, exc.message)

    return report_id


def _transition_report(report_id, actor, next_status, note):
    report = models.get_report(report_id, include_json=False)
    if not report:
        raise ServiceError("Report not found", 404)

    current_status = report["status"]
    allowed_next = ALLOWED_TRANSITIONS.get(current_status, set())
    if next_status not in allowed_next:
        raise ServiceError(
            f"Invalid report transition: {current_status} -> {next_status}",
            400,
        )

    models.update_report_status(report_id, next_status)
    models.add_report_tracking(report_id, next_status, actor["id"], note)
    return models.get_report_detail(report_id, include_json=True)


def _normalize_submission(data, actor, existing_report=None):
    student = dict(data.get("student") or {})
    phone = str(student.get("phone") or "").strip() or None
    name = _format_name_words(str(student.get("name") or "").strip())

    if not name:
        raise ServiceError("Student name is required", 400)
    if not phone and existing_report:
        phone = _effective_student_phone(existing_report)

    sales_id, sales_mode, sales_name = _normalize_sales_selection(student.get("sales_id"))

    course_type = str(student.get("course_type") or "").strip()
    course = str(student.get("course") or "").strip()
    if course_type not in COURSE_TYPES:
        raise ServiceError("Course must be General English, IELTS, SAT, or Other", 400)
    if course_type == "other":
        if not course:
            raise ServiceError("Course name is required when Other is selected", 400)
    else:
        course = COURSE_TYPES[course_type]

    class_type = str(student.get("class_type") or "").strip()
    if class_type not in CLASS_TYPES:
        raise ServiceError("Class type is required", 400)

    level = str(student.get("level") or "").strip()
    if course_type != "general_english":
        level = ""

    branch = actor.get("branch")
    student["phone"] = phone
    student["name"] = name
    student["branch"] = branch
    student["course"] = course
    student["course_type"] = course_type
    student["class_type"] = class_type
    student["level"] = level
    student["email"] = str(student.get("email") or "").strip() or None
    student["sales_id"] = sales_id
    student["sales_mode"] = sales_mode
    if sales_name:
        student["sales_name"] = sales_name
    else:
        student.pop("sales_name", None)
    student["teacher"] = _display_person_name(actor["name"])

    performance = dict(data.get("performance") or {})
    exam_score = str(performance.get("exam_score") or "").strip()
    if not exam_score:
        raise ServiceError("Exam score is required", 400)
    try:
        exam_score_value = float(exam_score)
    except ValueError as exc:
        raise ServiceError("Exam score must be a number", 400) from exc
    if exam_score_value < 0 or exam_score_value > 100:
        raise ServiceError("Exam score must be between 0 and 100", 400)
    performance["exam_score"] = (
        str(int(exam_score_value))
        if exam_score_value.is_integer()
        else str(exam_score_value).rstrip("0").rstrip(".")
    )

    report_type = (
        data.get("report_type")
        or student.get("report_type")
        or (existing_report or {}).get("report_type")
        or "midterm"
    )
    if report_type not in {"midterm", "final"}:
        raise ServiceError("Report type must be midterm or final", 400)

    normalized = dict(data)
    normalized["student"] = student
    normalized["performance"] = performance
    normalized["report_type"] = report_type
    return normalized


def _compose_report_payload(submission, ai_payload):
    payload = dict(ai_payload or {})
    payload["student"] = dict(submission.get("student") or {})
    payload["report_type"] = submission.get("report_type")
    payload["performance"] = dict(submission.get("performance") or {})
    payload["skills"] = dict(submission.get("skills") or {})
    payload["input_sections"] = dict(submission.get("sections") or {})
    return payload


def _generate_ai_report(data):
    student = data.get("student", {})
    name = student["name"]

    prompt = f"""
You are a senior academic English instructor writing a professional student progress report.

STRICT RULES:
- DO NOT invent names
- DO NOT use "the student"
- Professional academic tone
- Maintain professional academic tone
- No repetition
- No bullet points
- Provide clear analysis of strengths and weaknesses
- No simple sentences
- Fully developed paragraphs
- Use ONLY provided data
- Use this exact name: {name}
- Every sentence must end with a single period.
- Do not use ellipses or three-dot endings.
- No hallucination
- Return valid JSON only

Return JSON:
{{
  "student": {json.dumps(student, ensure_ascii=False)},
  "report_type": "{data['report_type']}",
  "sections": [
    {{"title": "INTRODUCTION", "content": "Generated paragraph."}},
    {{"title": "ATTENDANCE & PERFORMANCE", "content": "Generated paragraph."}},
    {{"title": "COMMUNICATION SKILLS", "content": "Generated paragraph."}},
    {{"title": "VOCABULARY DEVELOPMENT", "content": "Generated paragraph."}},
    {{"title": "GRAMMAR ANALYSIS", "content": "Generated paragraph."}},
    {{"title": "READING & WRITING", "content": "Generated paragraph."}},
    {{"title": "STRENGTHS", "content": "Generated paragraph."}},
    {{"title": "AREAS FOR IMPROVEMENT", "content": "Generated paragraph."}},
    {{"title": "FINAL RECOMMENDATION", "content": "Generated paragraph."}}
  ]
}}

DATA:
{json.dumps(data, ensure_ascii=False)}
"""

    try:
        raw = ai_client.generate_structured_text(prompt)
        report = json.loads(raw[raw.find("{") : raw.rfind("}") + 1])
    except Exception as exc:
        current_app.logger.warning("Remote text generation failed; using local report fallback: %s", exc)
        return _generate_local_report(data)

    for section in report.get("sections", []):
        content = section.get("content", "")
        content = content.replace("the student", name).replace("The student", name)
        section["content"] = _normalize_report_content(content)

    report["student"] = dict(student)
    report["report_type"] = data["report_type"]
    return report


def _generate_local_report(data):
    student = data.get("student", {})
    name = str(student.get("name") or "Student").strip()
    report_type = data.get("report_type") or "midterm"
    course = str(student.get("course") or "the selected course").strip()
    class_type = str(student.get("class_type") or "the assigned class type").strip()
    performance = data.get("performance") or {}
    skills = data.get("skills") or {}
    input_sections = data.get("sections") or {}

    skill_scores = _collect_skill_scores(data)
    snapshot_scores = _collect_snapshot_scores(data)
    strongest_skill, weakest_skill = _get_strength_markers(skill_scores)
    overall_score = _compute_average(list(skill_scores.values()) + list(snapshot_scores.values()))

    attendance = _score_phrase(performance.get("attendance"))
    assignment = _score_phrase(performance.get("assignment"))
    exam_score = str(performance.get("exam_score") or "").strip()
    communication_notes = _join_notes(input_sections.get("communication"))
    vocabulary_notes = _join_notes(input_sections.get("vocabulary"))
    grammar_notes = _join_notes(input_sections.get("grammar"))
    strengths = _join_notes(input_sections.get("strengths"))
    weaknesses = _join_notes(input_sections.get("weaknesses"))
    final_note = str(input_sections.get("final_note") or "").strip()

    sections = [
        {
            "title": "INTRODUCTION",
            "content": (
                f"{name} has completed the current {_format_report_type(report_type).lower()} cycle "
                f"for {course} in a {class_type} class with an overall performance average of {overall_score}/10. "
                "The report reflects classroom performance, submitted work, and the core language skill scores provided by the teacher."
            ),
        },
        {
            "title": "ATTENDANCE & PERFORMANCE",
            "content": (
                f"{name}'s attendance is currently {attendance}, assignment completion is {assignment}, "
                f"and the exam score is {exam_score}/100. "
                "These indicators should be monitored together because consistent attendance and regular "
                "assignment practice strongly support measurable language progress."
            ),
        },
        {
            "title": "COMMUNICATION SKILLS",
            "content": (
                communication_notes
                or f"{name}'s strongest measured area is {strongest_skill['label']}, while {weakest_skill['label']} "
                "requires the most focused support during speaking and classroom communication tasks."
            ),
        },
        {
            "title": "VOCABULARY DEVELOPMENT",
            "content": (
                vocabulary_notes
                or f"{name} should continue expanding vocabulary range and accuracy through topic-based practice, "
                "guided review, and frequent use of new words in complete spoken and written responses."
            ),
        },
        {
            "title": "GRAMMAR ANALYSIS",
            "content": (
                grammar_notes
                or f"{name}'s grammar development should focus on accuracy, sentence control, and consistent use "
                "of target structures across classroom activities and written assignments."
            ),
        },
        {
            "title": "READING & WRITING",
            "content": (
                f"In reading, {name} scored {_score_value(skills.get('reading'))}/10, and in writing, "
                f"{name} scored {_score_value(skills.get('writing'))}/10. Continued practice should connect "
                "reading comprehension with clear written production so progress remains balanced."
            ),
        },
        {
            "title": "STRENGTHS",
            "content": strengths or f"{name}'s strongest current area is {strongest_skill['label']}, which should be maintained through regular practice.",
        },
        {
            "title": "AREAS FOR IMPROVEMENT",
            "content": weaknesses or f"{name} should prioritize {weakest_skill['label']} with targeted support and short, consistent practice tasks.",
        },
        {
            "title": "FINAL RECOMMENDATION",
            "content": final_note or f"{name} should continue with a structured study plan that reinforces weaker skills while preserving current strengths.",
        },
    ]

    for section in sections:
        section["content"] = _normalize_report_content(section.get("content", ""))

    return {
        "student": dict(student),
        "report_type": report_type,
        "sections": sections,
    }


def _score_value(value):
    return _to_int(value)


def _score_phrase(value):
    score = _score_value(value)
    if score >= 8:
        return f"strong at {score}/10"
    if score >= 6:
        return f"developing at {score}/10"
    if score > 0:
        return f"in need of support at {score}/10"
    return "not yet scored"


def _join_notes(items):
    if isinstance(items, str):
        return items.strip()
    if not isinstance(items, list):
        return ""
    notes = [str(item or "").strip() for item in items]
    return " ".join(note for note in notes if note)


def _normalize_report_content(content):
    text = str(content or "").strip()
    if not text:
        return ""

    text = text.replace("…", ".")
    text = re.sub(r"\.{2,}", ".", text)
    text = re.sub(r"[!?](?=\s|$)", ".", text)
    text = re.sub(r"\s+", " ", text).strip()
    return _ensure_period(text)


def _display_person_name(name):
    clean_name = str(name or "").strip()
    if not clean_name:
        return ""
    if clean_name.lower().startswith(("mr. ", "ms. ")):
        prefix, rest = clean_name.split(" ", 1)
        return f"{prefix[:1].upper()}{prefix[1:].lower()} {_format_name_words(rest)}"
    first_name = clean_name.split()[0].lower()
    prefix = "Mr." if first_name in MR_NAMES else "Ms."
    return f"{prefix} {_format_name_words(clean_name)}"


def _format_name_words(name):
    return " ".join(_format_name_token(part) for part in str(name or "").split())


def _format_name_token(token):
    return "-".join(piece[:1].upper() + piece[1:].lower() if piece else "" for piece in token.split("-"))


def _ensure_period(text):
    text = str(text or "").strip()
    if not text:
        return ""
    text = text.rstrip(" .!?…")
    return f"{text}."


def generate_pdf_bytes(data):
    html = render_report_html(data, preview_mode=False)
    return convert_html_to_pdf(html)


def generate_docx_bytes(data):
    Document, WD_ALIGN_PARAGRAPH, Inches, Pt, RGBColor = _load_docx_dependencies()
    buffer = io.BytesIO()
    document = Document()
    student = data.get("student", {})
    report_type_label = _format_report_type(data.get("report_type"))
    issued_on = time.strftime("%d %B %Y")
    skill_scores = _collect_skill_scores(data)
    snapshot_scores = _collect_snapshot_scores(data)
    overall_score = _compute_average(list(skill_scores.values()) + list(snapshot_scores.values()))
    strongest_skill, weakest_skill = _get_strength_markers(skill_scores)

    _configure_document(document, Inches, Pt)
    _configure_header_footer(document.sections[0], Inches, Pt, issued_on)

    _add_cover_page(
        document,
        student,
        report_type_label,
        issued_on,
        overall_score,
        strongest_skill,
        weakest_skill,
        data,
        WD_ALIGN_PARAGRAPH,
    )

    document.add_page_break()
    _add_performance_page(document, data, overall_score, strongest_skill, weakest_skill, Inches)

    document.add_page_break()
    _add_analysis_page(document, data)

    document.add_page_break()
    _add_action_plan_page(document, data, strongest_skill, weakest_skill)

    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def convert_docx_to_pdf(docx_bytes):
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise ServiceError("PDF conversion dependencies are not installed", 500) from exc

    pythoncom.CoInitialize()

    base_temp_dir = os.path.join(os.getcwd(), ".runtime_tmp")
    os.makedirs(base_temp_dir, exist_ok=True)
    docx_file = tempfile.NamedTemporaryFile(prefix="report_", suffix=".docx", dir=base_temp_dir, delete=False)
    pdf_file = tempfile.NamedTemporaryFile(prefix="report_", suffix=".pdf", dir=base_temp_dir, delete=False)
    docx_path = docx_file.name
    pdf_path = pdf_file.name
    docx_file.close()
    pdf_file.close()
    try:
        with open(docx_path, "wb") as handle:
            handle.write(docx_bytes)

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False

        try:
            document = word.Documents.Open(docx_path)
            time.sleep(1)
            document.SaveAs(pdf_path, FileFormat=17)
            document.Close()
        except Exception as exc:
            word.Quit()
            raise ServiceError(f"PDF conversion failed: {exc}", 500) from exc

        word.Quit()

        if not os.path.exists(pdf_path):
            raise ServiceError("PDF generation failed", 500)

        with open(pdf_path, "rb") as handle:
            return handle.read()
    finally:
        for path in (docx_path, pdf_path):
            try:
                if os.path.exists(path):
                    os.remove(path)
            except OSError:
                pass


def convert_html_to_pdf(html):
    try:
        from playwright.sync_api import sync_playwright
    except ImportError as exc:
        raise ServiceError(
            "Playwright is not installed. Run 'pip install playwright' and then 'playwright install chromium'.",
            500,
        ) from exc

    try:
        with sync_playwright() as playwright:
            browser = playwright.chromium.launch()
            page = browser.new_page(viewport={"width": 1280, "height": 1810}, device_scale_factor=1.25)
            page.set_content(html, wait_until="load")
            pdf_bytes = page.pdf(
                format="A4",
                print_background=True,
                margin={"top": "12mm", "right": "12mm", "bottom": "14mm", "left": "12mm"},
                prefer_css_page_size=True,
            )
            browser.close()
            return pdf_bytes
    except Exception as exc:
        raise ServiceError(
            f"HTML to PDF generation failed: {exc}. Make sure Playwright Chromium is installed with 'playwright install chromium'.",
            500,
        ) from exc


def create_chart_bytes(data):
    plt = _load_matplotlib()
    buffer = io.BytesIO()
    skills = _collect_skill_scores(data)
    labels = list(skills.keys())
    values = list(skills.values())
    colors = ["#0F4C81", "#1D6F8F", "#2A9D8F", "#5AA9A7", "#84C5BE", "#B7D9D1"]

    figure, axis = plt.subplots(figsize=(6.9, 3.85))
    figure.patch.set_facecolor("#F7FAFC")
    axis.set_facecolor("#F7FAFC")
    bars = axis.bar(labels, values, color=colors[: len(labels)], width=0.62)
    axis.set_ylim(0, 10.6)
    axis.set_yticks(range(0, 11, 2))
    axis.grid(axis="y", color="#D9E2EC", linestyle="--", linewidth=0.8, alpha=0.8)
    axis.set_axisbelow(True)
    axis.spines["top"].set_visible(False)
    axis.spines["right"].set_visible(False)
    axis.spines["left"].set_color("#CBD5E1")
    axis.spines["bottom"].set_color("#CBD5E1")
    axis.tick_params(axis="x", rotation=18, labelsize=8.5, colors="#334155")
    axis.tick_params(axis="y", labelsize=9, colors="#475569")
    axis.set_title("Core Skill Performance", fontsize=13, fontweight="bold", color="#0F172A", pad=12)

    for bar, value in zip(bars, values):
        axis.text(
            bar.get_x() + (bar.get_width() / 2),
            value + 0.18,
            f"{value}",
            ha="center",
            va="bottom",
            fontsize=9,
            fontweight="bold",
            color="#0F172A",
        )

    figure.subplots_adjust(left=0.08, right=0.98, top=0.84, bottom=0.22)
    figure.savefig(buffer, format="png", dpi=220, bbox_inches="tight", pad_inches=0.15, facecolor=figure.get_facecolor())
    plt.close(figure)

    buffer.seek(0)
    return buffer.read()


def _store_pdf(report_id, pdf_bytes, student, report_type):
    provider = current_app.config.get("STORAGE_PROVIDER", "database").lower()
    storage_key = _build_storage_key(report_id, student, report_type)

    if provider == "s3":
        public_url = _store_pdf_s3(report_id, storage_key, pdf_bytes)
    else:
        public_url = _store_pdf_database(report_id, storage_key, pdf_bytes)

    models.update_report_pdf_url(report_id, public_url)
    return public_url


def _store_pdf_database(report_id, storage_key, pdf_bytes):
    models.save_report_document(
        report_id=report_id,
        storage_provider="database",
        storage_key=storage_key,
        content=pdf_bytes,
    )
    return _build_public_pdf_url(report_id)


def _store_pdf_s3(report_id, storage_key, pdf_bytes):
    try:
        import boto3
    except ImportError:
        return _store_pdf_database(report_id, storage_key, pdf_bytes)

    bucket_name = current_app.config.get("S3_BUCKET_NAME")
    if not bucket_name:
        return _store_pdf_database(report_id, storage_key, pdf_bytes)

    client = boto3.client(
        "s3",
        endpoint_url=current_app.config.get("S3_ENDPOINT_URL"),
        aws_access_key_id=current_app.config.get("AWS_ACCESS_KEY_ID"),
        aws_secret_access_key=current_app.config.get("AWS_SECRET_ACCESS_KEY"),
        region_name=current_app.config.get("AWS_REGION"),
    )
    client.put_object(
        Bucket=bucket_name,
        Key=storage_key,
        Body=pdf_bytes,
        ContentType="application/pdf",
        ACL="public-read",
    )

    models.save_report_document(
        report_id=report_id,
        storage_provider="s3",
        storage_key=storage_key,
        content=None,
    )

    custom_public_base = current_app.config.get("S3_PUBLIC_BASE_URL")
    if custom_public_base:
        return urljoin(custom_public_base.rstrip("/") + "/", storage_key)

    endpoint = current_app.config.get("S3_ENDPOINT_URL")
    if endpoint:
        return f"{endpoint.rstrip('/')}/{bucket_name}/{storage_key}"
    return f"https://{bucket_name}.s3.amazonaws.com/{storage_key}"


def _build_public_pdf_url(report_id):
    base_url = current_app.config.get("PUBLIC_BASE_URL")
    if not base_url:
        if has_request_context():
            base_url = request.url_root.rstrip("/")
        else:
            base_url = "http://localhost:5000"
    return f"{base_url.rstrip('/')}/storage/reports/{report_id}/pdf"


def _append_dataset_entry(input_data, output_data):
    entry = {"input": input_data, "output": output_data}
    try:
        if os.path.exists(DATASET_FILE):
            with open(DATASET_FILE, "r", encoding="utf-8") as handle:
                dataset = json.load(handle)
        else:
            dataset = []

        dataset.append(entry)
        with open(DATASET_FILE, "w", encoding="utf-8") as handle:
            json.dump(dataset, handle, ensure_ascii=False, indent=2)
    except Exception:
        pass


def _serialize_listing(row):
    report = dict(row)
    report["sent_student_email"] = bool(report.get("sent_student_email"))
    report["created_at"] = _format_display_timestamp(report.get("created_at"))
    report["sent_at"] = _format_display_timestamp(report.get("sent_at"))
    raw_student_id = str(report.get("student_id") or "").strip()
    display_phone = str(report.get("display_student_phone") or "").strip()
    if display_phone:
        report["student_id"] = display_phone
    elif raw_student_id.startswith("legacy-"):
        report["student_id"] = "-"
    else:
        report["student_id"] = raw_student_id or "-"
    if report.get("sales_name") == DIRECT_SALES_LABEL:
        report["sales_name"] = DIRECT_SALES_LABEL
    for key in ("student_name", "teacher_name", "created_by_name"):
        if report.get(key):
            report[key] = _format_name_words(report[key])
    if report.get("sales_name") and report["sales_name"] != DIRECT_SALES_LABEL:
        report["sales_name"] = _format_name_words(report["sales_name"])
    return report


def _can_view_report(actor, report):
    if actor.get("role") == "superadmin":
        return True
    if user_has_role(actor, "manager"):
        return True
    if user_has_role(actor, "admin", "operation", "sales_admin"):
        if report.get("status") == "draft":
            return False
        return report.get("student_branch") == actor.get("branch") or report.get("teacher_branch") == actor.get("branch")
    if user_has_role(actor, "teacher") and actor["id"] == report["teacher_id"]:
        return True
    if user_has_role(actor, "sales") and actor["id"] == report.get("sales_id"):
        return True
    return False


def _to_int(value):
    try:
        return int(str(value).strip())
    except Exception:
        return 0


def _normalize_id(value):
    if value in ("", None):
        return None
    try:
        return int(str(value).strip())
    except Exception:
        return None


def _normalize_sales_selection(value):
    raw_value = str(value).strip() if value is not None else ""
    if not raw_value:
        return None, None, None
    if raw_value.lower() == DIRECT_SALES_VALUE:
        return None, DIRECT_SALES_VALUE, DIRECT_SALES_LABEL
    if raw_value.isdigit():
        return int(raw_value), None, None
    return None, None, None


def _effective_student_phone(report):
    if not report:
        return None
    display_phone = str(report.get("display_student_phone") or "").strip()
    if display_phone:
        return display_phone
    student_id = str(report.get("student_id") or "").strip()
    if student_id and not student_id.startswith("legacy-"):
        return student_id
    report_payload = report.get("report_json") or {}
    if isinstance(report_payload, str):
        try:
            report_payload = json.loads(report_payload)
        except json.JSONDecodeError:
            report_payload = {}
    return str((report_payload.get("student") or {}).get("phone") or "").strip() or None


def _format_display_timestamp(value):
    raw = str(value or "").strip()
    if not raw:
        return "-"

    normalized = raw.replace("T", " ")
    for fmt in ("%Y-%m-%d %H:%M:%S.%f", "%Y-%m-%d %H:%M:%S"):
        try:
            parsed = datetime.strptime(normalized, fmt).replace(tzinfo=timezone.utc)
            return parsed.astimezone(DISPLAY_TIMEZONE).strftime("%Y-%m-%d %H:%M:%S")
        except ValueError:
            continue
    return raw


def _build_storage_key(report_id, student, report_type):
    student_phone = re.sub(r"[^0-9A-Za-z]+", "", str(student.get("phone") or "pending"))
    student_name = _slugify(student.get("name") or "student")
    report_kind = _slugify(report_type or "report")
    return f"reports/{student_phone}/{student_phone}-{student_name}-{report_kind}-{report_id}.pdf"


def _slugify(value):
    value = re.sub(r"[^0-9A-Za-z]+", "-", str(value).strip().lower())
    return value.strip("-") or "student"


def _configure_document(document, Inches, Pt):
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(10.5)


def _configure_header_footer(section, Inches, Pt, issued_on):
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(8.0))
    header_table.columns[0].width = Inches(2.0)
    header_table.columns[1].width = Inches(6.0)

    logo_cell = header_table.cell(0, 0)
    info_cell = header_table.cell(0, 1)
    _set_cell_shading(logo_cell, "0F2D52")
    _set_cell_shading(info_cell, "0F2D52")

    logo_path = "static/logo.png"
    logo_paragraph = logo_cell.paragraphs[0]
    if os.path.exists(logo_path):
        logo_paragraph.add_run().add_picture(logo_path, width=Inches(1.05))
    else:
        _append_colored_text(logo_paragraph, "iEnglish", 11, True, "FFFFFF")

    info_paragraph = info_cell.paragraphs[0]
    _append_colored_text(info_paragraph, "iEnglish\n", 12, True, "FFFFFF")
    _append_colored_text(info_paragraph, f"Issued {issued_on}", 8.5, False, "D6E3F1")

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = 1
    _append_colored_text(footer_paragraph, "Page ", 8.5, False, "64748B")
    _append_page_number(footer_paragraph)


def _add_cover_page(document, student, report_type_label, issued_on, overall_score, strongest_skill, weakest_skill, data, align):
    hero = document.add_table(rows=1, cols=1)
    hero_cell = hero.cell(0, 0)
    _set_cell_shading(hero_cell, "0F2D52")
    hero_paragraph = hero_cell.paragraphs[0]
    hero_paragraph.alignment = align.CENTER
    _append_colored_text(hero_paragraph, "STUDENT PROGRESS REPORT\n", 22, True, "FFFFFF")
    _append_colored_text(hero_paragraph, f"{report_type_label} | {student.get('name', 'Student')}", 12, False, "D6E3F1")

    subtitle = document.add_paragraph()
    subtitle.alignment = align.CENTER
    _append_colored_text(
        subtitle,
        "A polished academic report designed for families, academic operations, and PDF archiving.",
        10.5,
        False,
        "475569",
    )

    meta_table = document.add_table(rows=2, cols=3)
    meta_entries = [
        ("Student", student.get("name") or "-"),
        ("Teacher", student.get("teacher") or "-"),
        ("Course", student.get("course") or "-"),
        ("Level", student.get("level") or "-"),
        ("Class Type", student.get("class_type") or "-"),
        ("Issue Date", issued_on),
    ]
    for index, (label, value) in enumerate(meta_entries):
        cell = meta_table.cell(index // 3, index % 3)
        _set_cell_shading(cell, "F8FBFF")
        paragraph = cell.paragraphs[0]
        _append_colored_text(paragraph, f"{label}\n", 9, True, "0F4C81")
        _append_colored_text(paragraph, str(value), 11, False, "0F172A")

    document.add_paragraph("")

    card_table = document.add_table(rows=1, cols=4)
    card_specs = [
        ("Overall Score", f"{overall_score}/10", "0F4C81"),
        ("Strongest Skill", strongest_skill["label"], "1D6F8F"),
        ("Support Focus", weakest_skill["label"], "C97A40"),
        ("Report Type", report_type_label, "2A9D8F"),
    ]
    for idx, (title, value, color) in enumerate(card_specs):
        cell = card_table.cell(0, idx)
        _set_cell_shading(cell, color)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align.CENTER
        _append_colored_text(paragraph, f"{title}\n", 9, True, "EAF2F8")
        _append_colored_text(paragraph, str(value), 13, True, "FFFFFF")

    document.add_paragraph("")

    summary_table = document.add_table(rows=1, cols=2)
    left_cell = summary_table.cell(0, 0)
    right_cell = summary_table.cell(0, 1)
    _set_cell_shading(left_cell, "FCFDFE")
    _set_cell_shading(right_cell, "FCFDFE")

    summary_paragraph = left_cell.paragraphs[0]
    _append_colored_text(summary_paragraph, "Executive Summary\n", 13, True, "0F2D52")
    _append_colored_text(
        summary_paragraph,
        _build_executive_summary(data, student, overall_score, strongest_skill, weakest_skill),
        10.5,
        False,
        "334155",
    )

    highlights = [
        ("Performance Band", _score_band(overall_score)),
        ("Attendance", f"{_collect_snapshot_scores(data)['Attendance']}/10"),
        ("Assignments", f"{_collect_snapshot_scores(data)['Assignment']}/10"),
        ("Recommendation", _shorten_text(_extract_section_content(data, "FINAL RECOMMENDATION"), 110)),
    ]
    for title, value in highlights:
        paragraph = right_cell.add_paragraph()
        _append_colored_text(paragraph, f"{title}: ", 10, True, "0F4C81")
        _append_colored_text(paragraph, value or "-", 10, False, "334155")


def _add_performance_page(document, data, overall_score, strongest_skill, weakest_skill, Inches):
    heading = document.add_paragraph()
    _append_colored_text(heading, "Performance Dashboard", 18, True, "0F2D52")

    intro = document.add_paragraph()
    _append_colored_text(
        intro,
        "This page visualizes core language performance together with attendance and assignment completion in a layout that stays clear in both PDF and print.",
        10.5,
        False,
        "475569",
    )

    chart_table = document.add_table(rows=1, cols=2)
    chart_table.columns[0].width = Inches(5.15)
    chart_table.columns[1].width = Inches(2.45)
    left_cell = chart_table.cell(0, 0)
    right_cell = chart_table.cell(0, 1)

    try:
        left_cell.paragraphs[0].add_run().add_picture(io.BytesIO(create_chart_bytes(data)), width=Inches(4.9))
    except Exception:
        left_cell.text = "Core skill chart unavailable."

    try:
        right_cell.paragraphs[0].add_run().add_picture(io.BytesIO(create_progress_chart_bytes(data)), width=Inches(2.25))
    except Exception:
        right_cell.text = "Snapshot chart unavailable."

    document.add_paragraph("")

    dashboard = document.add_table(rows=2, cols=3)
    insights = [
        ("Overall Standing", f"{overall_score}/10", _score_band(overall_score)),
        ("Strongest Area", strongest_skill["label"], f"{strongest_skill['value']}/10"),
        ("Growth Priority", weakest_skill["label"], f"{weakest_skill['value']}/10"),
        ("Teacher Observation", _shorten_text(_extract_section_content(data, "COMMUNICATION SKILLS"), 95), ""),
        ("Academic Highlight", _shorten_text(_extract_section_content(data, "STRENGTHS"), 95), ""),
        ("Intervention Focus", _shorten_text(_extract_section_content(data, "AREAS FOR IMPROVEMENT"), 95), ""),
    ]
    for idx, (title, value, extra) in enumerate(insights):
        cell = dashboard.cell(idx // 3, idx % 3)
        _set_cell_shading(cell, "F8FBFF")
        paragraph = cell.paragraphs[0]
        _append_colored_text(paragraph, f"{title}\n", 9, True, "0F4C81")
        _append_colored_text(paragraph, value, 10.5, False, "0F172A")
        if extra:
            _append_colored_text(paragraph, f"\n{extra}", 9, False, "64748B")


def _add_analysis_page(document, data):
    heading = document.add_paragraph()
    _append_colored_text(heading, "Detailed Teacher Analysis", 18, True, "0F2D52")

    sections_to_render = [
        ("INTRODUCTION", "0F4C81"),
        ("ATTENDANCE & PERFORMANCE", "1D6F8F"),
        ("COMMUNICATION SKILLS", "2A9D8F"),
        ("VOCABULARY DEVELOPMENT", "4C956C"),
        ("GRAMMAR ANALYSIS", "7A8F2B"),
        ("READING & WRITING", "9A6C2F"),
    ]
    for title, color in sections_to_render:
        content = _extract_section_content(data, title)
        if content:
            _add_section_panel(document, title, content, color)


def _add_action_plan_page(document, data, strongest_skill, weakest_skill):
    heading = document.add_paragraph()
    _append_colored_text(heading, "Action Plan & Recommendation", 18, True, "0F2D52")

    two_col = document.add_table(rows=1, cols=2)
    strengths_cell = two_col.cell(0, 0)
    focus_cell = two_col.cell(0, 1)
    _set_cell_shading(strengths_cell, "F4FBF8")
    _set_cell_shading(focus_cell, "FFF8F2")

    strengths_paragraph = strengths_cell.paragraphs[0]
    _append_colored_text(strengths_paragraph, "Strengths to Maintain\n", 12, True, "1C6E56")
    _append_colored_text(
        strengths_paragraph,
        _extract_section_content(data, "STRENGTHS") or f"{strongest_skill['label']} is currently the student's strongest area.",
        10.5,
        False,
        "334155",
    )

    focus_paragraph = focus_cell.paragraphs[0]
    _append_colored_text(focus_paragraph, "Areas for Improvement\n", 12, True, "B45309")
    _append_colored_text(
        focus_paragraph,
        _extract_section_content(data, "AREAS FOR IMPROVEMENT") or f"{weakest_skill['label']} should remain the immediate improvement focus.",
        10.5,
        False,
        "334155",
    )

    document.add_paragraph("")

    recommendation_panel = document.add_table(rows=1, cols=1)
    recommendation_cell = recommendation_panel.cell(0, 0)
    _set_cell_shading(recommendation_cell, "F8FBFF")
    recommendation_paragraph = recommendation_cell.paragraphs[0]
    _append_colored_text(recommendation_paragraph, "Final Recommendation\n", 13, True, "0F2D52")
    _append_colored_text(
        recommendation_paragraph,
        _extract_section_content(data, "FINAL RECOMMENDATION"),
        10.5,
        False,
        "334155",
    )

    document.add_paragraph("")
    next_steps_heading = document.add_paragraph()
    _append_colored_text(next_steps_heading, "Recommended Next Steps", 13, True, "0F2D52")
    for step in _derive_next_steps(data, strongest_skill, weakest_skill):
        paragraph = document.add_paragraph(style="List Bullet")
        _append_colored_text(paragraph, step, 10.2, False, "334155")

    signoff = document.add_paragraph()
    _append_colored_text(signoff, "Prepared by Academic CRM reporting workflow", 9.5, True, "64748B")


def create_progress_chart_bytes(data):
    plt = _load_matplotlib()
    buffer = io.BytesIO()
    snapshot = _collect_snapshot_scores(data)
    labels = list(snapshot.keys())
    values = list(snapshot.values())
    colors = ["#0F4C81", "#2A9D8F"]

    figure, axes = plt.subplots(1, 2, figsize=(3.55, 2.05))
    figure.patch.set_facecolor("#F7FAFC")
    for axis, label, value, color in zip(axes, labels, values, colors):
        axis.set_facecolor("#F7FAFC")
        axis.pie(
            [value, max(10 - value, 0.001)],
            startangle=90,
            colors=[color, "#E2E8F0"],
            counterclock=False,
            wedgeprops={"width": 0.28, "edgecolor": "#F7FAFC"},
        )
        axis.text(0, 0.05, f"{value}/10", ha="center", va="center", fontsize=10, fontweight="bold", color="#0F172A")
        axis.text(0, -0.23, label, ha="center", va="center", fontsize=8, color="#475569")
        axis.set_aspect("equal")

    figure.subplots_adjust(left=0.04, right=0.96, top=0.94, bottom=0.1, wspace=0.12)
    figure.savefig(buffer, format="png", dpi=220, bbox_inches="tight", pad_inches=0.1, facecolor=figure.get_facecolor())
    plt.close(figure)
    buffer.seek(0)
    return buffer.read()


def _collect_skill_scores(data):
    skills = data.get("skills", {})
    return {
        "Reading": _to_int(skills.get("reading")),
        "Writing": _to_int(skills.get("writing")),
        "Listening": _to_int(skills.get("listening")),
        "Speaking": _to_int(skills.get("speaking")),
        "Vocabulary": _to_int(skills.get("vocab")),
        "Grammar": _to_int(skills.get("grammar")),
    }


def _collect_snapshot_scores(data):
    performance = data.get("performance", {})
    return {
        "Attendance": _to_int(performance.get("attendance")),
        "Assignment": _to_int(performance.get("assignment")),
    }


def _compute_average(values):
    values = [value for value in values if value is not None]
    if not values:
        return 0
    return round(sum(values) / len(values), 1)


def _get_strength_markers(skill_scores):
    items = list(skill_scores.items()) or [("Reading", 0)]
    strongest = max(items, key=lambda item: item[1])
    weakest = min(items, key=lambda item: item[1])
    return {"label": strongest[0], "value": strongest[1]}, {"label": weakest[0], "value": weakest[1]}


def _format_report_type(report_type):
    return "Final Report" if report_type == "final" else "Midterm Report"


def _build_executive_summary(data, student, overall_score, strongest_skill, weakest_skill):
    introduction = _extract_section_content(data, "INTRODUCTION")
    if introduction:
        lead = _shorten_text(introduction, 210)
    else:
        lead = (
            f"{student.get('name', 'The student')} is currently performing at an overall average of {overall_score}/10, "
            f"with particular strength in {strongest_skill['label']} and the clearest growth need in {weakest_skill['label']}."
        )
    return lead


def _score_band(score):
    if score >= 9:
        return "Outstanding"
    if score >= 7.5:
        return "Strong"
    if score >= 6:
        return "Developing Well"
    if score >= 4:
        return "Needs Reinforcement"
    return "Priority Support"


def _extract_section_content(data, title):
    normalized = str(title or "").strip().upper()
    for section in data.get("sections", []):
        if str(section.get("title", "")).strip().upper() == normalized:
            return str(section.get("content") or "").strip()
    return ""


def _shorten_text(text, limit):
    text = " ".join(str(text or "").split())
    if len(text) <= limit:
        return _normalize_report_content(text)
    trimmed = text[: limit - 1].rsplit(" ", 1)[0]
    return _ensure_period(trimmed)


def _derive_next_steps(data, strongest_skill, weakest_skill):
    steps = [
        f"Maintain {strongest_skill['label']} through consistent weekly practice so current strengths remain stable.",
        f"Prioritize {weakest_skill['label']} with targeted classroom support and short homework cycles.",
    ]
    focus = _extract_section_content(data, "AREAS FOR IMPROVEMENT")
    if focus:
        steps.append(_shorten_text(focus, 135))
    recommendation = _extract_section_content(data, "FINAL RECOMMENDATION")
    if recommendation:
        steps.append(_shorten_text(recommendation, 140))
    while len(steps) < 4:
        steps.append("Review progress at the next reporting checkpoint and adjust the support plan based on measurable classroom performance.")
    return steps[:4]


def _add_section_panel(document, title, content, color):
    panel = document.add_table(rows=1, cols=1)
    cell = panel.cell(0, 0)
    _set_cell_shading(cell, "FCFDFE")
    paragraph = cell.paragraphs[0]
    _append_colored_text(paragraph, f"{title}\n", 12, True, color)
    _append_colored_text(paragraph, content, 10.5, False, "334155")
    document.add_paragraph("")


def _set_cell_shading(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell_properties.append(shading)


def _append_colored_text(paragraph, text, size, bold, color):
    _, _, _, Pt, RGBColor = _load_docx_dependencies()
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    rgb = color.lstrip("#")
    run.font.color.rgb = RGBColor.from_string(rgb)
    return run


def _append_page_number(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _to_data_uri(content, mime_type):
    return f"data:{mime_type};base64,{base64.b64encode(content).decode('ascii')}"


def _load_logo_data_uri():
    logo_path = os.path.join(os.getcwd(), "static", "logo.png")
    if not os.path.exists(logo_path):
        return ""
    try:
        with open(logo_path, "rb") as handle:
            return _to_data_uri(handle.read(), "image/png")
    except OSError:
        return ""




def _load_docx_dependencies():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise ServiceError("The python-docx package is not installed", 500) from exc

    return Document, WD_ALIGN_PARAGRAPH, Inches, Pt, RGBColor


def _load_matplotlib():
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError as exc:
        raise ServiceError("The matplotlib package is not installed", 500) from exc

    return plt
